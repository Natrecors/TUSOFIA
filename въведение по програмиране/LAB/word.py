from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Създаване на документ
document = Document()

# 1. Настройки на страницата (A4, полета 2 см)
section = document.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# 8. Горен колонтитул
header = section.header
header_para = header.paragraphs[0]
header_para.text = "Иван Иванов, 10 клас, номер 5"
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 7. Изображение (Трябва файл pic.jpg в папката)
# Бележка: Позиционирането "горе вляво" с обтичане е сложно за скрипт,
# добавяме я в началото, потребителят трябва да я премести ръчно ако е нужно обтичане.
try:
    document.add_picture('pic.jpg', width=Cm(5), height=Cm(5))
except:
    print("Внимание: Снимката 'pic.jpg' липсва и не беше добавена.")

# 2. Заглавие
title = document.add_paragraph("Влечуги")
title_fmt = title.paragraph_format
title_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_fmt.space_before = Pt(12)
title_fmt.space_after = Pt(12)
run = title.runs[0]
run.font.name = 'Georgia'
run.font.size = Pt(24)
run.font.bold = True

# Текст за първите абзаци
body_texts = [
    "Влечугите (Reptilia) са клас гръбначни животни, включват дишащи атмосферен въздух четирикраки, обикновено покрити с люспи.","Макар и все още използвана в повечето съвременни системи за класификация, групата на влечугите е парафилетична, като обединява всички Амниота с изключение на птиците, бозайниците и техни изчезнали роднини.","Влечугите са животни с непостоянна телесна температура (пойкилотермни) – не произвеждат достатъчно топлина, за да поддържат постоянна телесна температура, а вместо това разчитат на топлообмен с околната среда, за да регулират вътрешната си температура.", "Повечето влечуги са месоядни и яйцеснасящи, някои са яйцеживораждащи, а други са същински живораждащи.","Влечугите възникват преди около 320 – 310 милиона години по време на периода на карбон. Еволюират от напреднали подобни на земноводни влечуги, които стават все по-адаптирани към живот на сушата. Много групи влечуги изчезват, включително динозаври, птерозаври и водните групи като ихтиозаврите. Съвременните представители са разпространени на всички континенти, с изключение на Антарктида, като най-голямо разпространение имат в тропиците и субтропиците. Обитават подземни хабитати, пустини, гори, блата, океани и други. Размерите им варират от дължина 1,6 cm при гекона Sphaerodactylus ariasae до дължина 6 m и маса над 1 000 kg при соленоводния крокодил (Crocodylus porosus)."
]

# 3. Форматиране на абзаците
for text in body_texts:
    p = document.add_paragraph(text)
    p_fmt = p.paragraph_format
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_fmt.space_after = Pt(6)
    p_fmt.first_line_indent = Cm(1.5)
    p_fmt.line_spacing = 1.5
    run = p.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(12)

# 4. Абзац "Съвременните влечуги..."
sub_p = document.add_paragraph("Съвременните влечуги се разделят на 4 разреда:")
sub_fmt = sub_p.paragraph_format
sub_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
sub_fmt.space_after = Pt(12)
sub_fmt.left_indent = Cm(0.5)
run = sub_p.runs[0]
run.font.name = 'Arial'
run.font.size = Pt(12)
run.font.underline = True

# 5. Списък
bullets = [
    "Testudines (Костенурки) – около 300 вида;",
    "Squamata (Люспести) – включва змии, гущери и други – около 7 900 вида;",
    "Rhynchocephalia (Клюноглави) – 2 вида;",
    "Crocodylia (Крокодили) – 23 вида."
]

for item in bullets:
    p = document.add_paragraph(item, style='List Bullet')
    p_fmt = p.paragraph_format
    p_fmt.left_indent = Cm(1.5) # Общ отстъп
    # Hanging се контролира автоматично от стила List Bullet, но тук го форсираме
    p_fmt.first_line_indent = Cm(-0.63) # Стандартно за bullet
    p_fmt.space_before = Pt(6)
    p_fmt.space_after = Pt(6)
    run = p.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.italic = True

# 6. Последен абзац
last_text = "Науката изучаваща земноводни и влечуги се нарича херпетология (на гръцки: ἑρπετόν [херпетон] – „пълзящ“; и λογία [-логия] – „слово“, „наука“)."
p = document.add_paragraph(last_text)
p_fmt = p.paragraph_format
p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_fmt.space_after = Pt(6)
p_fmt.first_line_indent = Cm(1.5)
p_fmt.line_spacing = 1.5
run = p.runs[0]
run.font.name = 'Arial'
run.font.size = Pt(12)
run.font.bold = True

document.save('vlechugi_klas_nomer.docx')
print("Файлът е създаден успешно!")